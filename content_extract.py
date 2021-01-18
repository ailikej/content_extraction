import json
import argparse
import sys
from docx import Document

def getAllParagraphsText(document):
    paragraphsText = [para.text for para in document.paragraphs]
    return '\n'.join(paragraphsText)


def getTableText(document):
	tableText = {}
	tableTextList = []
	chapterSummariesTable = (document.tables)[0]
	

	# convert table context from all cells into a list
	for row in chapterSummariesTable.rows:
		for cell in row.cells:
			tableTextList.append(cell.text)

	# generate organized data logically chapter by chapter
	numOfColumns = len(chapterSummariesTable.rows[0].cells)
	numOfCells = len(tableTextList)
	for i in range(numOfColumns, numOfCells, numOfColumns):
		tableText[tableTextList[i]] = {'Summary': '; '.join(tableTextList[i+1].split('\n')), 'Rating': tableTextList[i+2]}
	return tableText

if __name__ == "__main__":

	# setup parser for input and output files path
	parser = argparse.ArgumentParser(description='Process content extraction from a Career Overview docx file to json document')
	parser.add_argument('input', type=str, help='A pre-formatted Career Overview docx file, example: CareerOverviewCandidateJohnDoe.docx')
	args = parser.parse_args()

	# create Document object by loading input docx file and error handdling for invalid input
	if args.input.split('.')[-1] != 'docx':
		sys.exit('Input file is not a docx file. Please try again')   

	try:
		document = Document(args.input)
	except:
		sys.exit('File does not exist. Please try again')


	# extract data for paragraphs sections 'General Summary' and 'Decision' from the Document object
	allParagraphsText = getAllParagraphsText(document)
	generalSummaryData = allParagraphsText.partition("\n\nGeneral Summary\n\n")[2].partition("\n\nChapter Summaries")[0]
	decisionData = allParagraphsText.partition("\n\nDecision\n\n")[2]

	# extract data for table section 'Chapter Summaries'
	chapterSummariesData = getTableText(document)

	# assemble all sections data in json format
	data = {}	
	data['General Summary'] = generalSummaryData
	data['Decision'] = decisionData
	data['Chapter Summaries'] = chapterSummariesData

	# save and output a json file for the result
	with open(args.input.partition(".docx")[0]+'.json', 'w') as outfile:
		json.dump(data,outfile)

	### testing ###
	# print(json.dumps(data, indent=4, sort_keys=False))

